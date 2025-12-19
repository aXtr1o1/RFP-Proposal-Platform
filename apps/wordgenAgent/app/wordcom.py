import json
import logging
from pathlib import Path
from typing import Tuple, Optional, List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from apps.wordgenAgent.app.config_setting import build_updated_config

logger = logging.getLogger("wordcom")

# ---------------------------------------------------------
# Alignment codes used internally
# ---------------------------------------------------------
WD_ALIGN_LEFT = 0
WD_ALIGN_CENTER = 1
WD_ALIGN_RIGHT = 2
WD_ALIGN_JUSTIFY = 3

# RTL/LTR flags
WD_READINGORDER_LTR = 0
WD_READINGORDER_RTL = 1

# ---------------------------------------------------------
# Default configuration
# ---------------------------------------------------------
default_CONFIG = {
    "visible_word": False,
    "output_path": "output/proposal.docx",
    "language_lcid": 1025,

    "default_alignment": WD_ALIGN_LEFT,
    "reading_order": WD_READINGORDER_LTR,

    "space_before": 0,
    "space_after": 6,
    "line_spacing_rule": 0,

    "orientation": 0,
    "margin_top": 72,
    "margin_bottom": 72,
    "margin_left": 72,
    "margin_right": 72,

    "table_autofit": True,
    "table_preferred_width": 100,

    "title_style": "Title",
    "heading_style": "Heading 1",
    "normal_style": "Normal",

    "font_size": 11,
    "heading_font_size": 14,
    "title_font_size": 16,
    "points_font_size": 11,
    "table_font_size": 10,

    "title_font_color": 0,
    "heading_font_color": 0,
    "content_font_color": 0,
    "table_font_color": 0,

    "table_border_visible": True,
    "table_border_color": 0,
    "table_border_line_style": 1,
    "table_border_line_width": 1,

    "table_header_shading_color": None,
    "table_body_shading_color": None,

    "enable_header": False,
    "enable_footer": False,
    "company_name": "",
    "company_tagline": "",
    "header_logo_path": "",
    "header_logo_width": 5,

    "footer_left_text": "",
    "footer_center_text": "",
    "footer_right_text": "",
    "footer_show_page_numbers": True,

    "bullet_char": "•",
}


# =========================
# Color Utilities
# =========================
def _bgr_int_to_rgb_tuple(bgr: Optional[int]) -> Tuple[int, int, int]:
    if bgr is None:
        return (0, 0, 0)
    return (
        (bgr >> 0) & 0xFF,
        (bgr >> 8) & 0xFF,
        (bgr >> 16) & 0xFF,
    )


def _bgr_int_to_hex(bgr: Optional[int]) -> Optional[str]:
    if bgr is None:
        return None
    r, g, b = _bgr_int_to_rgb_tuple(bgr)
    return f"{r:02X}{g:02X}{b:02X}"


# =========================
# Core Alignment + RTL patch - FIXED VERSION
# =========================
def _apply_alignment_and_direction(paragraph, code: int, rtl: bool):
    """
    Correct python-docx compatible RTL + alignment.
    Fixed to properly apply alignment regardless of RTL setting.
    """
    pf = paragraph.paragraph_format
    
    # FIXED: Always set alignment first, regardless of RTL
    if code == WD_ALIGN_RIGHT:
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif code == WD_ALIGN_CENTER:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif code == WD_ALIGN_JUSTIFY:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Get or create paragraph properties
    pPr = paragraph._p.get_or_add_pPr()

    # Set RTL properties only if RTL is enabled
    if rtl:
        # w:bidi → basic RTL support
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        bidi.set(qn("w:val"), "1")
        
        # w:rtl → right-to-left layout
        rtl_tag = pPr.find(qn("w:rtl"))
        if rtl_tag is None:
            rtl_tag = OxmlElement("w:rtl")
            pPr.append(rtl_tag)
        rtl_tag.set(qn("w:val"), "1")
    else:
        # Remove RTL properties if not RTL
        bidi = pPr.find(qn("w:bidi"))
        if bidi is not None:
            pPr.remove(bidi)
        
        rtl_tag = pPr.find(qn("w:rtl"))
        if rtl_tag is not None:
            pPr.remove(rtl_tag)
    
    # FIXED: Don't set textAlignment for alignment - it interferes
    # Only use the paragraph format alignment which we set above
    text_align = pPr.find(qn("w:textAlignment"))
    if text_align is not None:
        pPr.remove(text_align)
    
    # Force justification setting at XML level for better compatibility
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    
    # Map alignment code to Word's jc values
    if code == WD_ALIGN_RIGHT:
        jc.set(qn("w:val"), "right")
    elif code == WD_ALIGN_CENTER:
        jc.set(qn("w:val"), "center")
    elif code == WD_ALIGN_JUSTIFY:
        jc.set(qn("w:val"), "both")
    else:
        jc.set(qn("w:val"), "left")


# =========================
# Paragraph Builders
# =========================
def _add_para(doc, text, style_name, size, color_bgr, bold, align_code, rtl):
    p = doc.add_paragraph()

    if style_name:
        try:
            p.style = style_name
        except Exception as e:
            logger.warning(f"Failed to apply style '{style_name}': {e}")

    run = p.add_run(text or "")
    if size:
        run.font.size = Pt(size)

    if color_bgr is not None:
        r, g, b = _bgr_int_to_rgb_tuple(color_bgr)
        run.font.color.rgb = RGBColor(r, g, b)

    run.font.bold = bool(bold)

    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # FIXED: Apply alignment and direction with corrected function
    _apply_alignment_and_direction(p, align_code, rtl)
    return p


def _add_bullet_para(doc, text, cfg, align_code, rtl):
    p = doc.add_paragraph()

    style_name = cfg.get("normal_style", "Normal")
    try:
        p.style = style_name
    except Exception:
        pass

    bullet_char = cfg.get("bullet_char", "•")
    size = cfg.get("points_font_size", 11)
    color_bgr = cfg.get("content_font_color", 0)
    r, g, b = _bgr_int_to_rgb_tuple(color_bgr)

    run_bullet = p.add_run(f"{bullet_char} ")
    run_bullet.font.size = Pt(size)
    run_bullet.font.color.rgb = RGBColor(r, g, b)

    run_text = p.add_run(text or "")
    run_text.font.size = Pt(size)
    run_text.font.color.rgb = RGBColor(r, g, b)

    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)

    _apply_alignment_and_direction(p, align_code, rtl)
    return p


# =========================
# Tables
# =========================
def _set_table_width_pct(tbl, pct: int):
    pct = max(1, min(100, int(pct or 100)))
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(pct * 50))


def _shade_cell(cell, fill_hex: Optional[str]):
    if not fill_hex:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)


def _add_table(doc, headers, rows, cfg, align_code, rtl):
    headers = headers or []
    rows = rows or []
    if not headers and not rows:
        return

    num_cols = len(headers) if headers else len(rows[0])
    num_rows = len(rows) + (1 if headers else 0)

    tbl = doc.add_table(rows=num_rows, cols=num_cols)
    tbl.autofit = cfg.get("table_autofit", True)
    _set_table_width_pct(tbl, cfg.get("table_preferred_width", 100))

    font_rgb = _bgr_int_to_rgb_tuple(cfg.get("table_font_color", 0))
    header_fill = _bgr_int_to_hex(cfg.get("table_header_shading_color"))
    body_fill = _bgr_int_to_hex(cfg.get("table_body_shading_color"))
    font_size = cfg.get("table_font_size", 10)

    row_idx = 0

    if headers:
        for c, h in enumerate(headers):
            cell = tbl.cell(row_idx, c)
            cell.text = ""
            p = cell.paragraphs[0]

            run = p.add_run(str(h))
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*font_rgb)

            _apply_alignment_and_direction(p, align_code, rtl)
            _shade_cell(cell, header_fill)

        row_idx += 1

    for row in rows:
        for c, val in enumerate(row):
            cell = tbl.cell(row_idx, c)
            cell.text = ""
            p = cell.paragraphs[0]

            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*font_rgb)

            _apply_alignment_and_direction(p, align_code, rtl)
            _shade_cell(cell, body_fill)

        row_idx += 1


# =========================
# MAIN BUILDER - FIXED VERSION
# =========================
def build_word_from_proposal(proposal, user_config, output_path, lang, visible=False):
    if isinstance(proposal, str):
        proposal = json.loads(proposal or "{}")

    if not isinstance(user_config, dict):
        user_config = {}

    cfg = build_updated_config(default_CONFIG, user_config)

    # FIXED: Determine RTL setting
    reading_order = cfg.get("reading_order", WD_READINGORDER_LTR)
    rtl_from_cfg = (reading_order == WD_READINGORDER_RTL)
    rtl_from_lang = str(lang or "").lower().startswith("ar")
    rtl = rtl_from_cfg or rtl_from_lang

    # FIXED: Get alignment from config - it's already set by build_updated_config
    align_code = cfg.get("default_alignment", WD_ALIGN_LEFT)
    
    # Validate alignment code is in valid range
    if align_code not in (WD_ALIGN_LEFT, WD_ALIGN_CENTER, WD_ALIGN_RIGHT, WD_ALIGN_JUSTIFY):
        logger.warning(f"Invalid alignment code {align_code}, using RIGHT for RTL, LEFT otherwise")
        align_code = WD_ALIGN_RIGHT if rtl else WD_ALIGN_LEFT

    logger.info(f"[wordcom] Effective alignment={align_code} (0=left,1=center,2=right,3=justify), rtl={rtl}")

    template_path = cfg.get("template_path")
    try:
        if template_path:
            doc = Document(template_path)
        else:
            doc = Document()
    except Exception:
        doc = Document()

    # Title
    title = proposal.get("title", "")
    if title.strip():
        _add_para(
            doc,
            title.strip(),
            cfg.get("title_style", "Title"),
            cfg.get("title_font_size", 16),
            cfg.get("title_font_color", 0),
            True,
            align_code,
            rtl,
        )

    for sec in proposal.get("sections", []):
        heading = (sec.get("heading") or "").strip()
        content = (sec.get("content") or "").strip()
        points = sec.get("points") or []
        table = sec.get("table") or {}

        if heading:
            _add_para(
                doc,
                heading,
                cfg.get("heading_style", "Heading 1"),
                cfg.get("heading_font_size", 14),
                cfg.get("heading_font_color", 0),
                True,
                align_code,
                rtl,
            )

        if content:
            for line in content.split("\n"):
                t = line.strip()
                if t:
                    _add_para(
                        doc,
                        t,
                        cfg.get("normal_style", "Normal"),
                        cfg.get("font_size", 11),
                        cfg.get("content_font_color", 0),
                        False,
                        align_code,
                        rtl,
                    )

        for pt in points:
            t = (pt or "").strip()
            if t:
                _add_bullet_para(doc, t, cfg, align_code, rtl)

        headers = table.get("headers") or []
        rows = table.get("rows") or []
        if headers or rows:
            _add_table(doc, headers, rows, cfg, align_code, rtl)

    abs_out = str(Path(output_path).resolve())
    Path(abs_out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(abs_out)
    return abs_out